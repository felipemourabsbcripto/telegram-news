#!/usr/bin/env python3
"""
Gera documentação completa do Bot em formato DOCX
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    import subprocess
    subprocess.run(["pip3", "install", "python-docx"], check=True)
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE

from datetime import datetime

def create_documentation():
    doc = Document()
    
    # ==================== CAPA ====================
    title = doc.add_heading('🤖 Bot Telegram de Notícias Cripto', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Documentação Completa de Funcionalidades')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Versão 2.0 - Fevereiro 2026\n').bold = True
    info.add_run('Sistema automatizado de agregação, tradução e distribuição de notícias sobre criptomoedas')
    
    doc.add_page_break()
    
    # ==================== ÍNDICE ====================
    doc.add_heading('Índice', level=1)
    
    indice = [
        "1. Visão Geral do Sistema",
        "2. Painel Administrativo",
        "3. Fontes de Notícias",
        "4. Calendário de Eventos Cripto",
        "5. Inteligência Artificial (Groq)",
        "6. Formato das Postagens",
        "7. Temas e Categorias",
        "8. Horários de Postagem",
        "9. Analytics e Relatórios",
        "10. Banco de Dados",
        "11. Arquitetura Técnica",
        "12. Configurações",
    ]
    
    for item in indice:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # ==================== 1. VISÃO GERAL ====================
    doc.add_heading('1. Visão Geral do Sistema', level=1)
    
    doc.add_paragraph(
        'O Bot Telegram de Notícias Cripto é um sistema automatizado completo para '
        'agregação, processamento e distribuição de notícias sobre criptomoedas. '
        'O sistema coleta notícias de múltiplas fontes, traduz automaticamente para '
        'português, aplica filtros de relevância usando inteligência artificial e '
        'distribui para um canal do Telegram.'
    )
    
    doc.add_heading('Principais Recursos:', level=2)
    
    recursos = [
        "Agregação automática de notícias de 30+ fontes",
        "Tradução automática para português (Google Translate)",
        "Filtro de relevância com IA (Groq - Llama 3.1 70B)",
        "Resumo automático de notícias com IA",
        "Calendário de eventos cripto com alertas",
        "Painel administrativo com teclado inline",
        "Analytics e relatórios de desempenho",
        "Sistema de alertas para eventos importantes",
    ]
    
    for r in recursos:
        doc.add_paragraph(r, style='List Bullet')
    
    # ==================== 2. PAINEL ADMIN ====================
    doc.add_heading('2. Painel Administrativo', level=1)
    
    doc.add_paragraph(
        'O bot possui um painel de administração acessível via Telegram com menus '
        'interativos usando teclado inline (botões clicáveis).'
    )
    
    doc.add_heading('Comandos Disponíveis:', level=2)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Comando'
    hdr_cells[1].text = 'Descrição'
    
    comandos = [
        ('/start ou /config', 'Abre o painel de configuração principal'),
        ('/calendar ou /eventos', 'Acessa o calendário de eventos cripto'),
        ('/status', 'Mostra status atual do bot'),
        ('/help', 'Exibe mensagem de ajuda'),
    ]
    
    for i, (cmd, desc) in enumerate(comandos, 1):
        row = table.rows[i].cells
        row[0].text = cmd
        row[1].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('Menu Principal:', level=2)
    
    menus = [
        "📰 Fontes - Gerenciar fontes de notícias",
        "📅 Calendário Cripto - Eventos, conferências e discursos",
        "⏰ Horários - Configurar horários de postagem",
        "📝 Formato - Configurar formato das mensagens",
        "🏷️ Temas - Filtrar por categorias",
        "🤖 IA (Groq) - Configurações de inteligência artificial",
        "📊 Analytics - Relatórios e métricas",
        "▶️ Status - Ver status do sistema",
    ]
    
    for m in menus:
        doc.add_paragraph(m, style='List Bullet')
    
    # ==================== 3. FONTES ====================
    doc.add_heading('3. Fontes de Notícias', level=1)
    
    doc.add_paragraph(
        'O sistema suporta múltiplas fontes de notícias, divididas em fontes '
        'pré-configuradas (built-in) e fontes que podem ser adicionadas pelo usuário.'
    )
    
    doc.add_heading('3.1 Fontes Built-in:', level=2)
    
    fontes_builtin = [
        ("CoinDesk", "https://www.coindesk.com/"),
        ("CoinTelegraph", "https://cointelegraph.com/"),
        ("Decrypt", "https://decrypt.co/"),
        ("Bitcoin Magazine", "https://bitcoinmagazine.com/"),
        ("CryptoSlate", "https://cryptoslate.com/"),
        ("UToday", "https://u.today/"),
        ("Portal do Bitcoin", "https://portaldobitcoin.uol.com.br/"),
        ("CoinTelegraph BR", "https://br.cointelegraph.com/"),
        ("CriptoFácil", "https://www.criptofacil.com/"),
    ]
    
    table = doc.add_table(rows=len(fontes_builtin)+1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Fonte'
    table.rows[0].cells[1].text = 'URL'
    
    for i, (nome, url) in enumerate(fontes_builtin, 1):
        table.rows[i].cells[0].text = nome
        table.rows[i].cells[1].text = url
    
    doc.add_paragraph()
    
    doc.add_heading('3.2 Fontes Populares (Adição Rápida):', level=2)
    
    doc.add_heading('Internacionais:', level=3)
    
    internacionais = [
        "Whale Alert - Alertas de grandes transações",
        "Glassnode Insights - Análises on-chain",
        "TradingView News - Notícias de mercado",
        "The Block - Notícias institucionais",
        "Blockworks - Análises de mercado",
        "BeInCrypto - Notícias gerais",
        "Messari - Pesquisa e dados",
        "The Defiant - Foco em DeFi",
        "Daily Hodl - Notícias diárias",
        "CryptoPotato - Análises e notícias",
        "Binance Blog/News/Square - Ecossistema Binance",
    ]
    
    for f in internacionais:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_heading('Brasileiras:', level=3)
    
    brasileiras = [
        "Livecoins - Principal portal BR",
        "CriptoFácil - Notícias simplificadas",
        "Portal do Bitcoin - Foco em Bitcoin",
        "CoinTelegraph BR - Versão brasileira",
        "BeInCrypto BR - Versão brasileira",
        "InfoMoney Cripto - Seção de cripto",
        "Exame Future of Money - Análises",
        "Money Times Cripto - Mercado",
    ]
    
    for f in brasileiras:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_heading('Exchanges:', level=3)
    
    exchanges = [
        "Coinbase Blog",
        "Kraken Blog",
        "Mercado Bitcoin Blog",
    ]
    
    for f in exchanges:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_heading('3.3 Gerenciamento de Fontes:', level=2)
    
    gerenciamento = [
        "✅/❌ Toggle - Ativar ou desativar fonte",
        "🗑️ Excluir - Remover fonte completamente",
        "➕ Adicionar - Adicionar fonte personalizada com seletores CSS",
        "⭐ Fontes Populares - Menu com fontes pré-configuradas para adicionar",
    ]
    
    for g in gerenciamento:
        doc.add_paragraph(g, style='List Bullet')
    
    # ==================== 4. CALENDÁRIO ====================
    doc.add_heading('4. Calendário de Eventos Cripto', level=1)
    
    doc.add_paragraph(
        'Sistema completo de calendário para acompanhar eventos importantes do '
        'mercado de criptomoedas, incluindo conferências, discursos econômicos '
        'e lançamentos de projetos.'
    )
    
    doc.add_heading('4.1 Visualizações Disponíveis:', level=2)
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Menu'
    table.rows[0].cells[1].text = 'Descrição'
    
    views = [
        ('📅 Eventos Hoje', 'Eventos do dia atual'),
        ('📆 Próximos 7 Dias', 'Agenda semanal'),
        ('🗓️ Próximos 30 Dias', 'Agenda mensal'),
        ('🎤 Discursos Importantes', 'Reuniões Fed/FOMC, falas econômicas'),
        ('🎪 Conferências 2026', 'Grandes eventos do ano'),
        ('🚀 Lançamentos', 'Updates de protocolos'),
    ]
    
    for i, (menu, desc) in enumerate(views, 1):
        table.rows[i].cells[0].text = menu
        table.rows[i].cells[1].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('4.2 Eventos Pré-carregados 2026:', level=2)
    
    doc.add_heading('Conferências:', level=3)
    
    conferencias = [
        ("ETHDenver", "24 Fev - 02 Mar", "Denver, EUA"),
        ("Paris Blockchain Week", "07-11 Abr", "Paris, França"),
        ("NFT.NYC", "15-17 Abr", "New York, EUA"),
        ("Token2049 Dubai", "28-29 Abr", "Dubai, UAE"),
        ("Bitcoin 2026 Conference", "15-17 Mai", "Nashville, EUA"),
        ("Consensus Miami", "26-28 Mai", "Miami, EUA"),
        ("Web Summit Rio", "15-18 Jun", "Rio de Janeiro, Brasil"),
        ("Blockchain Rio", "10-12 Ago", "Rio de Janeiro, Brasil"),
        ("Token2049 Singapore", "14-15 Set", "Singapura"),
        ("Gramado Summit", "20-22 Set", "Gramado, Brasil"),
        ("Devcon", "20-23 Out", "TBA"),
        ("Consensus Hong Kong", "10-12 Nov", "Hong Kong"),
        ("Blockchain Life Dubai", "08-10 Dez", "Dubai, UAE"),
    ]
    
    table = doc.add_table(rows=len(conferencias)+1, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Evento'
    table.rows[0].cells[1].text = 'Data'
    table.rows[0].cells[2].text = 'Local'
    
    for i, (evento, data, local) in enumerate(conferencias, 1):
        table.rows[i].cells[0].text = evento
        table.rows[i].cells[1].text = data
        table.rows[i].cells[2].text = local
    
    doc.add_paragraph()
    
    doc.add_heading('Discursos/Econômicos:', level=3)
    
    economicos = [
        ("FOMC Meeting - Fed", "28 Jan, 18 Mar, 06 Mai, 17 Jun, 29 Jul, 16 Set, 04 Nov, 16 Dez"),
        ("World Economic Forum Davos", "19-23 Jan"),
        ("Jackson Hole Symposium", "27-29 Ago"),
        ("G20 Summit", "21-22 Nov"),
    ]
    
    for evento, data in economicos:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{evento}: ").bold = True
        p.add_run(data)
    
    doc.add_heading('4.3 Sistema de Alertas:', level=2)
    
    alertas = [
        ("🔔 Alertas Ativos", "Liga/desliga todos os alertas"),
        ("⏰ 1 Dia Antes", "Alerta enviado 24 horas antes do evento"),
        ("🔔 1 Hora Antes", "Alerta enviado 1 hora antes do evento"),
        ("🎪 Alertar Conferências", "Filtro para conferências"),
        ("🎤 Alertar Discursos", "Filtro para discursos"),
        ("🚀 Alertar Lançamentos", "Filtro para lançamentos"),
    ]
    
    table = doc.add_table(rows=len(alertas)+1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Configuração'
    table.rows[0].cells[1].text = 'Descrição'
    
    for i, (config, desc) in enumerate(alertas, 1):
        table.rows[i].cells[0].text = config
        table.rows[i].cells[1].text = desc
    
    # ==================== 5. IA ====================
    doc.add_heading('5. Inteligência Artificial (Groq)', level=1)
    
    doc.add_paragraph(
        'O sistema utiliza a API do Groq com o modelo Llama 3.1 70B para '
        'processamento inteligente das notícias. O Groq oferece uma API gratuita '
        'com alta velocidade de resposta.'
    )
    
    doc.add_heading('5.1 Funcionalidades de IA:', level=2)
    
    ia_funcs = [
        ("🧠 Resumir com IA", "Gera resumos concisos e informativos das notícias em português"),
        ("🎯 Filtrar Relevância", "Avalia cada notícia com nota de 1 a 10, descartando as irrelevantes"),
        ("😎 Adicionar Emojis", "Adiciona emojis contextuais aos títulos das notícias"),
        ("🏷️ Classificar Tema", "Categoriza automaticamente em: news, analysis, whale, defi, nft, etc."),
    ]
    
    table = doc.add_table(rows=len(ia_funcs)+1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Funcionalidade'
    table.rows[0].cells[1].text = 'Descrição'
    
    for i, (func, desc) in enumerate(ia_funcs, 1):
        table.rows[i].cells[0].text = func
        table.rows[i].cells[1].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('5.2 Configuração:', level=2)
    
    config_ia = [
        "API Principal: Groq (gratuita)",
        "Modelo: Llama 3.1 70B Versatile",
        "API Backup: OpenAI GPT-3.5 Turbo (opcional)",
        "Nota mínima de relevância configurável (padrão: 5)",
    ]
    
    for c in config_ia:
        doc.add_paragraph(c, style='List Bullet')
    
    # ==================== 6. FORMATO ====================
    doc.add_heading('6. Formato das Postagens', level=1)
    
    doc.add_heading('6.1 Opções de Conteúdo:', level=2)
    
    formato = [
        ("🔗 Mostrar Link", "Incluir link da notícia original"),
        ("🖼️ Mostrar Imagem", "Incluir imagem da notícia (se disponível)"),
        ("📹 Mostrar Vídeo", "Incluir vídeo (se disponível)"),
        ("🌐 Traduzir", "Tradução automática para português"),
        ("🤖 Resumir com IA", "Usar IA para gerar resumo"),
    ]
    
    table = doc.add_table(rows=len(formato)+1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Opção'
    table.rows[0].cells[1].text = 'Descrição'
    
    for i, (opt, desc) in enumerate(formato, 1):
        table.rows[i].cells[0].text = opt
        table.rows[i].cells[1].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('6.2 Estilos de Postagem:', level=2)
    
    estilos = [
        ("📄 Completo", "Título + conteúdo completo da notícia"),
        ("📋 Resumido", "Título + até 300 caracteres do conteúdo"),
        ("📌 Só Título", "Apenas o título da notícia"),
    ]
    
    table = doc.add_table(rows=len(estilos)+1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Estilo'
    table.rows[0].cells[1].text = 'Descrição'
    
    for i, (estilo, desc) in enumerate(estilos, 1):
        table.rows[i].cells[0].text = estilo
        table.rows[i].cells[1].text = desc
    
    # ==================== 7. TEMAS ====================
    doc.add_heading('7. Temas e Categorias', level=1)
    
    doc.add_paragraph(
        'As notícias são automaticamente classificadas em categorias, permitindo '
        'filtrar o tipo de conteúdo que será postado no canal.'
    )
    
    temas = [
        ("📰 News", "Notícias gerais do mercado cripto"),
        ("📊 Analysis", "Análises de mercado e preço"),
        ("🔗 On-Chain", "Dados on-chain e métricas de blockchain"),
        ("🐋 Whale", "Movimentação de grandes investidores (baleias)"),
        ("💥 Liquidation", "Notícias sobre liquidações no mercado"),
        ("🏦 Exchange", "Notícias de exchanges e corretoras"),
        ("⚖️ Regulation", "Regulamentação e legislação"),
        ("🌾 DeFi", "Finanças descentralizadas"),
        ("🎨 NFT", "NFTs e metaverso"),
    ]
    
    table = doc.add_table(rows=len(temas)+1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Tema'
    table.rows[0].cells[1].text = 'Descrição'
    
    for i, (tema, desc) in enumerate(temas, 1):
        table.rows[i].cells[0].text = tema
        table.rows[i].cells[1].text = desc
    
    # ==================== 8. HORÁRIOS ====================
    doc.add_heading('8. Horários de Postagem', level=1)
    
    doc.add_paragraph(
        'É possível configurar horários específicos para postagem de notícias, '
        'definindo o tema e quantidade máxima de posts para cada horário.'
    )
    
    doc.add_heading('Formato de Configuração:', level=2)
    
    p = doc.add_paragraph()
    p.add_run('HH:MM tema quantidade').bold = True
    
    doc.add_paragraph('Exemplos:', style='List Bullet')
    doc.add_paragraph('09:00 news 5 - Posta até 5 notícias às 9h', style='List Bullet 2')
    doc.add_paragraph('14:00 analysis 3 - Posta até 3 análises às 14h', style='List Bullet 2')
    doc.add_paragraph('20:00 whale 2 - Posta até 2 alertas de baleias às 20h', style='List Bullet 2')
    
    # ==================== 9. ANALYTICS ====================
    doc.add_heading('9. Analytics e Relatórios', level=1)
    
    doc.add_paragraph(
        'O sistema coleta métricas de todas as postagens enviadas, permitindo '
        'acompanhar o desempenho do canal.'
    )
    
    doc.add_heading('Relatórios Disponíveis:', level=2)
    
    relatorios = [
        ("📈 Relatório Hoje", "Posts enviados, visualizações, forwards e reações do dia"),
        ("📊 Relatório Semanal", "Métricas consolidadas dos últimos 7 dias"),
        ("🏆 Top 10 Posts", "Posts com mais visualizações"),
        ("📰 Por Fonte", "Performance de cada fonte de notícias"),
        ("🏷️ Por Tema", "Performance por categoria/tema"),
        ("🔄 Atualizar Métricas", "Refresh manual dos dados"),
    ]
    
    table = doc.add_table(rows=len(relatorios)+1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Relatório'
    table.rows[0].cells[1].text = 'Descrição'
    
    for i, (rel, desc) in enumerate(relatorios, 1):
        table.rows[i].cells[0].text = rel
        table.rows[i].cells[1].text = desc
    
    # ==================== 10. BANCO DE DADOS ====================
    doc.add_heading('10. Banco de Dados', level=1)
    
    doc.add_paragraph(
        'O sistema utiliza PostgreSQL para armazenamento persistente de '
        'configurações, notícias postadas, eventos e métricas.'
    )
    
    doc.add_heading('Tabelas:', level=2)
    
    tabelas = [
        ("bot_config", "Configurações do bot em formato JSON"),
        ("scheduled_posts", "Horários agendados de postagem"),
        ("post_analytics", "Métricas de cada post (views, forwards, etc.)"),
        ("crypto_events", "Eventos do calendário cripto"),
        ("news_*", "Uma tabela por fonte para controle de duplicatas"),
    ]
    
    table = doc.add_table(rows=len(tabelas)+1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Tabela'
    table.rows[0].cells[1].text = 'Descrição'
    
    for i, (tab, desc) in enumerate(tabelas, 1):
        table.rows[i].cells[0].text = tab
        table.rows[i].cells[1].text = desc
    
    # ==================== 11. ARQUITETURA ====================
    doc.add_heading('11. Arquitetura Técnica', level=1)
    
    doc.add_heading('11.1 Fluxo de Funcionamento:', level=2)
    
    fluxo = [
        "1. Busca notícias de cada fonte ativa",
        "2. Verifica duplicatas no banco de dados",
        "3. Traduz para português (se ativado)",
        "4. Filtra por relevância com IA (se ativado)",
        "5. Classifica tema automaticamente",
        "6. Resume com IA (se ativado)",
        "7. Adiciona emojis (se ativado)",
        "8. Envia para o canal Telegram",
        "9. Salva métricas no banco",
        "10. Aguarda intervalo e repete",
    ]
    
    for f in fluxo:
        doc.add_paragraph(f, style='List Number')
    
    doc.add_heading('11.2 Threads do Sistema:', level=2)
    
    threads = [
        ("Thread Principal", "Processa comandos do painel admin"),
        ("Thread News Fetcher", "Busca e posta notícias (ciclo de 5 min)"),
        ("Thread Event Alerts", "Verifica e envia alertas de eventos (ciclo de 1h)"),
    ]
    
    table = doc.add_table(rows=len(threads)+1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Thread'
    table.rows[0].cells[1].text = 'Função'
    
    for i, (thread, func) in enumerate(threads, 1):
        table.rows[i].cells[0].text = thread
        table.rows[i].cells[1].text = func
    
    doc.add_heading('11.3 Tecnologias Utilizadas:', level=2)
    
    techs = [
        "Python 3.12",
        "PostgreSQL 15 (via Docker)",
        "SQLAlchemy (ORM)",
        "Telegram Bot API",
        "Groq API (Llama 3.1 70B)",
        "Google Translate (deep-translator)",
        "BeautifulSoup4 (web scraping)",
        "lxml (parsing HTML)",
    ]
    
    for t in techs:
        doc.add_paragraph(t, style='List Bullet')
    
    # ==================== 12. CONFIGURAÇÕES ====================
    doc.add_heading('12. Configurações', level=1)
    
    doc.add_heading('12.1 Variáveis de Ambiente:', level=2)
    
    vars_env = [
        ("TELEGRAM_TOKEN", "Token do bot Telegram"),
        ("CHANNEL_ID", "ID do canal para postagem"),
        ("DATABASE_URL", "URL de conexão PostgreSQL"),
        ("GROQ_API_KEY", "Chave da API Groq"),
        ("OPENAI_API_KEY", "Chave da API OpenAI (opcional)"),
        ("ADMIN_ID", "ID do admin (opcional)"),
    ]
    
    table = doc.add_table(rows=len(vars_env)+1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Variável'
    table.rows[0].cells[1].text = 'Descrição'
    
    for i, (var, desc) in enumerate(vars_env, 1):
        table.rows[i].cells[0].text = var
        table.rows[i].cells[1].text = desc
    
    doc.add_heading('12.2 Arquivos do Projeto:', level=2)
    
    arquivos = [
        ("admin_bot.py", "Bot principal com painel admin (~2000 linhas)"),
        ("run_bot.py", "Bot simples original"),
        ("docker-compose.yml", "Configuração do PostgreSQL"),
        ("requirements.txt", "Dependências Python"),
        ("gerar_documentacao.py", "Script para gerar este documento"),
    ]
    
    table = doc.add_table(rows=len(arquivos)+1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Arquivo'
    table.rows[0].cells[1].text = 'Descrição'
    
    for i, (arq, desc) in enumerate(arquivos, 1):
        table.rows[i].cells[0].text = arq
        table.rows[i].cells[1].text = desc
    
    doc.add_page_break()
    
    # ==================== RODAPÉ ====================
    doc.add_heading('Contato e Suporte', level=1)
    
    doc.add_paragraph(
        'Este documento foi gerado automaticamente em ' + 
        datetime.now().strftime('%d/%m/%Y às %H:%M') + '.'
    )
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Bot Telegram: ').bold = True
    p.add_run('@GornicapitalnesBot')
    
    p = doc.add_paragraph()
    p.add_run('Canal: ').bold = True
    p.add_run('ID -1003753019442')
    
    # Salvar
    filename = '/workspaces/telegram-news/Bot_Cripto_Documentacao.docx'
    doc.save(filename)
    print(f'✅ Documento salvo em: {filename}')
    return filename

if __name__ == '__main__':
    create_documentation()
